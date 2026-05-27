"""
core/forensic_reporter.py — Professional forensic Word report generator (v40.0).

Generates .docx incident reports using python-docx.
Output: logs/reports/incident_<ID>_<timestamp>.docx
"""

import asyncio
from datetime import datetime, timezone
from pathlib import Path
from loguru import logger

_REPORTS_DIR = Path("logs/reports")
_REPORTS_DIR.mkdir(parents=True, exist_ok=True)

_EXEC_SUMMARY_SYSTEM = """You are a senior cybersecurity consultant
writing an executive summary for a non-technical stakeholder.
Explain what happened, its impact, and key recommendations.
Maximum 3 paragraphs. Professional tone. No jargon."""


async def generate_forensic_report(
    incident: dict,
    related_episodes: list[dict],
    broadcast_fn,
    ollama_client,
    model: str,
    screenshot_paths: list[str] = [],
) -> str | None:
    """
    Generate a professional .docx forensic incident report.
    Returns file path or None on failure.
    """
    inc_id    = incident.get("incident_id", "UNK")
    severity  = incident.get("severity_score", 0)
    phase     = incident.get("kill_chain_phase", "Unknown")
    techniques= incident.get("mitre_techniques", [])
    hosts     = list(incident.get("involved_hosts", set()))
    pids      = list(incident.get("involved_pids", set()))
    sub_events= incident.get("sub_events", [])[:15]

    await broadcast_fn({
        "type":        "report_generating",
        "incident_id": inc_id,
        "format":      "docx",
        "timestamp":   datetime.now(timezone.utc).isoformat(),
    })

    exec_prompt = (
        f"Incident {inc_id}: severity {severity}/10, "
        f"phase={phase}, techniques={techniques}, "
        f"hosts={hosts}. "
        "Write executive summary for non-technical stakeholder."
    )
    try:
        resp = await asyncio.wait_for(
            ollama_client.chat.completions.create(
                model    = model,
                messages = [
                    {"role": "system", "content": _EXEC_SUMMARY_SYSTEM},
                    {"role": "user",   "content": exec_prompt},
                ],
                stream = False,
                extra_body = {"options": {"num_ctx": 1024, "temperature": 0.3}},
            ),
            timeout=30.0,
        )
        exec_summary = resp.choices[0].message.content.strip()
    except Exception:
        exec_summary = (
            f"A security incident (ID: {inc_id}) was detected with "
            f"severity score {severity:.1f}/10. The incident involved "
            f"{len(techniques)} MITRE ATT&CK techniques across "
            f"{len(hosts)} host(s)."
        )

    loop = asyncio.get_running_loop()
    doc_path = await loop.run_in_executor(
        None,
        _build_docx,
        incident, exec_summary, sub_events,
        techniques, hosts, pids, screenshot_paths,
        inc_id,
    )

    if doc_path:
        logger.info(f"FORENSIC_REPORTER: report saved → {Path(doc_path).name}")
        await broadcast_fn({
            "type":        "report_generated",
            "incident_id": inc_id,
            "format":      "docx",
            "filename":    Path(doc_path).name,
            "filepath":    doc_path,
            "timestamp":   datetime.now(timezone.utc).isoformat(),
        })

    return doc_path


def _build_docx(
    incident: dict,
    exec_summary: str,
    sub_events: list[dict],
    techniques: list[str],
    hosts: list,
    pids: list,
    screenshot_paths: list[str],
    inc_id: str,
) -> str | None:
    """Blocking — runs in executor. Builds the Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import hashlib

        doc = Document()

        section = doc.sections[0]
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = section.right_margin   = Inches(1)
        section.top_margin  = section.bottom_margin  = Inches(1)

        def _heading(text: str, level: int = 1) -> None:
            p = doc.add_heading(text, level=level)
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.color.rgb = RGBColor(0x00, 0xFF, 0x41)

        def _para(text: str, bold: bool = False,
                  color: str = None) -> None:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = bold
            if color:
                r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
                run.font.color.rgb = RGBColor(r, g, b)

        def _table_row(table, cells: list[str],
                        bold: bool = False) -> None:
            row = table.add_row()
            for i, cell_text in enumerate(cells):
                cell = row.cells[i]
                cell.text = str(cell_text)
                if bold:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

        # Cover page
        doc.add_paragraph()
        title = doc.add_heading("FORENSIC INCIDENT REPORT", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(f"Incident ID: {inc_id}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_p = doc.add_paragraph(
            f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}"
        )
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sev_color = "FF0000" if incident.get("severity_score",0) >= 8 else \
                    "FF6600" if incident.get("severity_score",0) >= 6 else "FFAA00"
        _para(
            f"SEVERITY: {incident.get('severity_score',0):.1f}/10 — "
            f"{incident.get('kill_chain_phase','Unknown')}",
            bold=True, color=sev_color,
        )

        doc.add_paragraph(
            "Generated by: JARVIS Purple Team Platform v40.0"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Executive Summary
        _heading("1. Executive Summary")
        doc.add_paragraph(exec_summary)

        # Incident Timeline
        _heading("2. Incident Timeline")
        if sub_events:
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = "Light Shading"
            _table_row(tbl, ["#", "Event Type", "Process/Host", "Technique"], bold=True)
            for i, evt in enumerate(sub_events, 1):
                _table_row(tbl, [
                    str(i),
                    evt.get("type", ""),
                    evt.get("process", evt.get("attacker_ip", "?")),
                    evt.get("technique", "—"),
                ])

        # Technical Analysis
        _heading("3. Technical Analysis")
        _para(f"Kill Chain Phase: {incident.get('kill_chain_phase','?')}", bold=True)
        _para(f"Involved Hosts: {', '.join(str(h) for h in hosts) or 'None identified'}")
        _para(f"Involved PIDs: {', '.join(str(p) for p in pids) or 'None identified'}")

        # IOC Summary
        _heading("4. Indicators of Compromise (IOCs)")
        ioc_tbl = doc.add_table(rows=1, cols=3)
        ioc_tbl.style = "Light Shading"
        _table_row(ioc_tbl, ["IOC Type", "Value", "Notes"], bold=True)
        for host in hosts[:10]:
            _table_row(ioc_tbl, ["IP Address", str(host), "Observed in incident"])
        for pid in pids[:5]:
            _table_row(ioc_tbl, ["Process ID", str(pid), "Suspicious PID"])

        # MITRE ATT&CK Mapping
        _heading("5. MITRE ATT&CK Mapping")
        att_tbl = doc.add_table(rows=1, cols=3)
        att_tbl.style = "Light Shading"
        _table_row(att_tbl, ["Technique ID", "Name", "ATT&CK URL"], bold=True)
        for tech in techniques[:15]:
            url = f"https://attack.mitre.org/techniques/{tech.replace('.','/')}"
            _table_row(att_tbl, [tech, f"ATT&CK {tech}", url])

        # Screenshots
        if screenshot_paths:
            _heading("6. Evidence Screenshots")
            for path_str in screenshot_paths[:5]:
                p = Path(path_str)
                if p.exists() and p.suffix.lower() in {".png", ".jpg"}:
                    try:
                        doc.add_picture(str(p), width=Inches(5.5))
                        doc.add_paragraph(
                            f"SHA-256: {hashlib.sha256(p.read_bytes()).hexdigest()}"
                        ).style = "Caption"
                    except Exception:
                        pass

        # Remediation
        _heading("7. Remediation Recommendations")
        recs = [
            "1. Isolate affected hosts immediately via network segmentation",
            "2. Collect volatile memory evidence before any remediation",
            "3. Revoke and rotate all credentials potentially exposed",
            "4. Apply relevant patches for exploited vulnerabilities",
            "5. Update detection rules based on observed IOCs",
            "6. Conduct full forensic analysis of affected systems",
            "7. Review and harden vulnerable configurations",
        ]
        for rec in recs:
            doc.add_paragraph(rec, style="List Bullet")

        # Save
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"incident_{inc_id}_{ts}.docx"
        filepath = _REPORTS_DIR / filename
        doc.save(str(filepath))
        return str(filepath)

    except Exception as e:
        logger.error(f"FORENSIC_REPORTER: docx build error: {e}")
        return None
